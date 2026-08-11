import io
import qrcode

from docx import Document as DocxDocument
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _generate_qr_bytes(data):
    """يولّد صورة QR في الذاكرة (PNG) بنفس ألوان هوية النظام"""
    qr = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_M, box_size=10, border=2)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#1a2b4c", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def stamp_docx_with_qr(input_path, output_path, verify_url, document_number):
    """
    يُدرج رمز QR والرقم المرجعي/التسلسلي داخل ترويسة (header) الصفحة الأولى
    فقط من ملف Word (.docx)، عبر خاصية "ترويسة مختلفة للصفحة الأولى"
    (Different First Page) الخاصة بـ Word، بحيث لا يظهر الرمز في بقية صفحات
    المستند ولا يُعاد رسمه أو تحويله (يبقى الملف .docx قابلاً للتحرير).

    يُطبَّق التعديل على القسم الأول (section) فقط من المستند، لأن الصفحة
    الأولى الفعلية للملف تنتمي دائمًا للقسم الأول بغض النظر عن عدد الأقسام.

    ملاحظة: تعمل هذه الدالة فقط مع صيغة .docx الحديثة (XML)، ولا تدعم صيغة
    .doc القديمة (الثنائية) التي لا تدعمها مكتبة python-docx إطلاقًا.
    """
    doc = DocxDocument(input_path)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.first_page_header

    # تفريغ الفقرة الافتراضية الأولى في الترويسة (إن وُجدت) قبل الإدراج
    if header.paragraphs and header.paragraphs[0].runs:
        for run in list(header.paragraphs[0].runs):
            run.text = ""
        qr_paragraph = header.paragraphs[0]
    elif header.paragraphs:
        qr_paragraph = header.paragraphs[0]
    else:
        qr_paragraph = header.add_paragraph()

    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_run = qr_paragraph.add_run()
    qr_run.add_picture(_generate_qr_bytes(verify_url), width=Cm(2.4))

    number_paragraph = header.add_paragraph()
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number_run = number_paragraph.add_run(f"الرقم المرجعي / التسلسلي: {document_number}")
    number_run.bold = True
    number_run.font.size = Pt(11)

    doc.save(output_path)
    return output_path
